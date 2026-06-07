"""Create minimal DOCX templates with {{ }} placeholders for P1 validation."""
from docx import Document
from pathlib import Path


def make_activity_plan():
    doc = Document()
    doc.add_heading("活动方案", 0)
    doc.add_paragraph("")
    doc.add_paragraph("活动主要内容：")
    doc.add_paragraph("{{ activity_content }}")
    doc.add_paragraph("")
    doc.add_paragraph("活动开始时间：{{ start_time }}")
    doc.add_paragraph("活动结束时间：{{ end_time }}")
    doc.add_paragraph("共计天数：{{ total_days }}")
    doc.add_paragraph("开幕式时间：{{ opening_start }} 至 {{ opening_end }}")
    doc.add_paragraph("")
    doc.add_paragraph("工作人员：{{ staff_count }}人")
    doc.add_paragraph("演员：{{ performer_count }}人")
    doc.add_paragraph("嘉宾：{{ guest_count }}人")
    doc.add_paragraph("主要活动日人数：{{ opening_crowd }}人")
    doc.add_paragraph("平日人数：{{ regular_crowd }}人")
    doc.add_paragraph("")
    doc.add_paragraph("搭建方案：")
    doc.add_paragraph("{{ construction_plan }}")
    doc.add_paragraph("")
    doc.add_paragraph("负责人联系方式：{{ contact_phone }}")
    doc.add_paragraph("备注：{{ remarks }}")
    return doc


def make_security_plan():
    doc = Document()
    doc.add_heading("安保方案", 0)
    doc.add_paragraph("")
    doc.add_paragraph("安保人员配置：")
    doc.add_paragraph("{{ security_staff_config }}")
    doc.add_paragraph("安保人员数量：{{ security_staff_count }}人")
    doc.add_paragraph("")
    doc.add_paragraph("动线设计：")
    doc.add_paragraph("{{ movement_plan }}")
    doc.add_paragraph("")
    doc.add_paragraph("安保设备清单：")
    doc.add_paragraph("{{ equipment_list }}")
    doc.add_paragraph("")
    doc.add_paragraph("应急预案：")
    doc.add_paragraph("{{ emergency_plan }}")
    doc.add_paragraph("")
    doc.add_paragraph("医疗救护措施：{{ medical_plan }}")
    doc.add_paragraph("消防措施：{{ fire_plan }}")
    doc.add_paragraph("人流管控方案：{{ crowd_control }}")
    doc.add_paragraph("")
    doc.add_paragraph("安保负责人审核签名：{{ manager_signature }}")
    return doc


def make_risk_assessment():
    doc = Document()
    doc.add_heading("风险评估报备表", 0)
    doc.add_paragraph("填报单位（盖章）：{{ reporting_unit }}")
    doc.add_paragraph("填报日期：{{ report_date }}")
    doc.add_paragraph("项目名称：{{ project_name }}")
    doc.add_paragraph("活动类型：{{ activity_type }}")
    doc.add_paragraph("活动名称：{{ activity_name }}")
    doc.add_paragraph("主办方：{{ sponsor }}")
    doc.add_paragraph("承办方：{{ organizer }}")
    doc.add_paragraph("活动参与方：{{ participants }}")
    doc.add_paragraph("活动时间：{{ activity_start }} 至 {{ activity_end }}")
    doc.add_paragraph("活动地点：{{ activity_location }}（{{ is_indoor }}）")
    doc.add_paragraph("场所类型：{{ location_type }}")
    doc.add_paragraph("活动内容：{{ activity_content }}")
    doc.add_paragraph("活动规模：{{ crowd_scale }}人")
    doc.add_paragraph("工作人员：{{ staff_count }}人")
    doc.add_paragraph("安保人员：{{ security_count }}人")
    doc.add_paragraph("销售门票：{{ has_tickets }}")
    doc.add_paragraph("媒体直播：{{ has_media }}")
    doc.add_paragraph("{% if media_name %}媒体名称：{{ media_name }}（{{ media_type }}）{% endif %}")
    doc.add_paragraph("")
    doc.add_paragraph("主要风险因素：")
    doc.add_paragraph("{% for rf in risk_factors %}· {{ rf }}\n{% endfor %}")
    doc.add_paragraph("")
    doc.add_paragraph("防范化解措施：")
    doc.add_paragraph("{% for mm in mitigation_measures %}· {{ mm }}\n{% endfor %}")
    doc.add_paragraph("")
    doc.add_paragraph("联系人：{{ contact_person }}")
    doc.add_paragraph("联系电话：{{ contact_phone }}")
    doc.add_paragraph("评估主体负责人签字：{{ assessor_signature }}")
    doc.add_paragraph("安保负责人审核签名：{{ manager_signature }}")
    return doc


def make_responsibility_letter():
    doc = Document()
    doc.add_heading("主办单位安全和消防责任确认书", 0)
    doc.add_paragraph("我单位（{{ sponsor_unit }}）在{{ venue_name }}举办活动的安全和消防责任，依法确认如下：")
    doc.add_paragraph("")
    items = [
        "一、已具体制订活动安全工作方案，明确各项安全措施：{{ security_measures_confirmed }}",
        "二、保证临时搭建和使用的设施安全，沒有安全隐患：{{ facilities_safe }}",
        "三、已按照负责许可的公安机关的要求，配备必要的安全检查设备：{{ security_check_equipped }}",
        "四、严格按照核准的活动场所容纳人员数量、划定的区域组织活动：{{ capacity_compliant }}",
        "五、已落实医疗救护、灭火、应急疏散等应急救援措施并组织演练：{{ emergency_measures_confirmed }}",
        "六、对妨碍活动安全的行为及时予以制止：{{ misconduct_response }}",
        "七、已配备与活动安全工作需要相适应的专业保安人员：{{ professional_security_staffed }}",
        "八、已为活动的安全工作提供必要的保障，临时用电按规范：{{ temporary_electricity_safe }}",
    ]
    for text in items:
        doc.add_paragraph(text)
    doc.add_paragraph("")
    doc.add_paragraph("活动主办单位（公章）：{{ sponsor_seal }}")
    doc.add_paragraph("活动安全负责人（签字）：{{ security_leader_signature }}")
    doc.add_paragraph("活动安全负责人：{{ security_leader_name }}")
    doc.add_paragraph("确认日期：{{ confirm_date }}")
    doc.add_paragraph("确认地点：{{ confirm_location }}")
    doc.add_paragraph("安保负责人审核签名：{{ manager_signature }}")
    return doc


def make_filing_commitment():
    doc = Document()
    doc.add_heading("备案承诺书", 0)
    doc.add_paragraph("{{ commitment_content }}")
    doc.add_paragraph("")
    doc.add_paragraph("申请人：{{ applicant_name }}")
    doc.add_paragraph("申请人签字：{{ applicant_signature }}")
    doc.add_paragraph("备案日期：{{ filing_date }}")
    return doc


TEMPLATES = {
    "activity_plan": make_activity_plan,
    "security_plan": make_security_plan,
    "risk_assessment": make_risk_assessment,
    "responsibility_letter": make_responsibility_letter,
    "filing_commitment": make_filing_commitment,
}

ROOT = Path(__file__).parent.parent / "app" / "templates"

for name, factory in TEMPLATES.items():
    path = ROOT / name / "template.docx"
    doc = factory()
    doc.save(str(path))
    print(f"Created {path}")
