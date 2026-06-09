"""
Rebuild all 5 template DOCX files from source government forms.

Strategy per template:
- activity_plan: Build new table from scratch (too many rows to delete from source)
- security_plan: Build new single-column table from scratch
- risk_assessment: Edit converted source DOCX, replace text with placeholders
- responsibility_letter: Edit converted source DOCX, replace with placeholders + set fonts
- filing_commitment: Edit existing template, set fonts + keep paragraph structure

Fonts: title=FZXiaoBiaoSong-B05S, label=KaiTi_GB2312, body=FangSong_GB2312
"""

import copy
import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt
from lxml import etree

TEMPLATES_ROOT = Path(__file__).resolve().parent.parent / "app" / "templates"

# Font names as registered in the system
FONT_TITLE = "FZXiaoBiaoSong-B05S"  # 方正小标宋简体
FONT_LABEL = "KaiTi_GB2312"  # 楷体_GB2312
FONT_BODY = "FangSong_GB2312"  # 仿宋_GB2312

# Source converted files
SRC_DIR = Path("/tmp")


def _set_font(run, font_name, size_pt=None, bold=False):
    """Set font on a run, preserving other properties."""
    run.font.name = font_name
    # Set East-Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold


def _set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))


def _add_bordered_table(doc, rows, cols):
    """Add a table with thin black borders on all cells."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    return table


def _add_cell_text(cell, text, font_name=FONT_BODY, size_pt=10.5, bold=False, alignment=None):
    """Set cell text with font. Clears existing paragraphs."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    _set_font(run, font_name, size_pt, bold)
    return run


def _clear_row(row):
    """Clear all cell text in a row (for unmapped rows)."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''


def _set_row_height(row, emu):
    """Set exact row height in EMU (1 cm = 360000 EMU)."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = etree.SubElement(tr, qn('w:trPr'))
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = etree.SubElement(trPr, qn('w:trHeight'))
    trHeight.set(qn('w:val'), str(emu))
    trHeight.set(qn('w:hRule'), 'atLeast')


def _reduce_cell_spacing(cell):
    """Remove top/bottom paragraph spacing inside a cell."""
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


# ---------------------------------------------------------------------------
# activity_plan — build new table from scratch
# ---------------------------------------------------------------------------
def build_activity_plan():
    """Build activity plan template. Simplified table, only mapped fields."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("活动方案")
    _set_font(title_run, FONT_TITLE, 22, bold=True)

    # Table: label | value layout
    # Row descriptions: (label, field_placeholder, label_width_ratio)
    fields = [
        ("活动名称", "{{ activity_name }}"),
        ("主办单位", "{{ sponsor }}"),
        ("活动负责人及联系方式（手机）", "{{ contact_phone }}"),
        ("活动时间", "{{ start_time }} 至 {{ end_time }}，共计 {{ total_days }} 天"),
        ("开幕式时间", "{% if has_opening == '是' %}{{ opening_start }} 至 {{ opening_end }}{% else %}无{% endif %}"),
        ("主要活动内容", "{{ activity_content }}"),
        ("工作人员数量", "{{ staff_count }}"),
        ("演员数量", "{{ performer_count }}"),
        ("嘉宾数量", "{{ guest_count }}"),
        ("主要活动日人数规模", "{{ opening_crowd }}"),
        ("平日人数规模", "{{ regular_crowd }}"),
        ("搭建方案（含材料明细、平面图、效果图、标注尺寸）", "{{ construction_plan }}"),
        ("备注", "备注：现场秩序维护及安全保障由主办方负责。\n{{ remarks }}"),
    ]

    table = _add_bordered_table(doc, len(fields), 2)

    for i, (label, value) in enumerate(fields):
        _add_cell_text(table.cell(i, 0), label, FONT_LABEL, 10.5, bold=True)
        if value:
            _add_cell_text(table.cell(i, 1), value, FONT_BODY, 10.5)

    # Remove unused paragraphs from default template
    for p in doc.paragraphs:
        if not p.text.strip() and p != title_p:
            p._element.getparent().remove(p._element)

    output_path = TEMPLATES_ROOT / "activity_plan" / "template.docx"
    doc.save(str(output_path))
    print(f"  ✅ activity_plan → {output_path}")


# ---------------------------------------------------------------------------
# security_plan — build new single-column wide table
# ---------------------------------------------------------------------------
def build_security_plan():
    """Build security plan template. Single column wide rows (label + content)."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("安保方案")
    _set_font(title_run, FONT_TITLE, 22, bold=True)

    fields = [
        ("安保人员配置", "{{ security_staff_config }}"),
        ("安保人员数量", "{{ security_staff_count }}"),
        ("人员疏导方案", "{{ movement_plan }}"),
        ("设备清单", "{{ equipment_list }}"),
        ("应急预案", "{{ emergency_plan }}"),
        ("医疗保障方案", "{% if risk_level == '高风险' %}{{ medical_plan }}{% endif %}"),
        ("消防方案", "{% if risk_level != '低风险' %}{{ fire_plan }}{% endif %}"),
        ("人群控制方案", "{% if risk_level == '高风险' %}{{ crowd_control }}{% endif %}"),
    ]

    table = _add_bordered_table(doc, len(fields), 1)

    for i, (label, value) in enumerate(fields):
        cell = table.cell(i, 0)
        # Clear default paragraph
        cell.paragraphs[0].clear()
        # Label paragraph
        label_p = cell.paragraphs[0]
        label_run = label_p.add_run(label)
        _set_font(label_run, FONT_LABEL, 10.5, bold=True)
        # Value paragraph (separate paragraph for docxtpl placeholders)
        if value:
            val_p = cell.add_paragraph()
            val_run = val_p.add_run(value)
            _set_font(val_run, FONT_BODY, 10.5)

    output_path = TEMPLATES_ROOT / "security_plan" / "template.docx"
    doc.save(str(output_path))
    print(f"  ✅ security_plan → {output_path}")


# ---------------------------------------------------------------------------
# risk_assessment — build from scratch, source-inspired layout
# ---------------------------------------------------------------------------
def build_risk_assessment():
    """Build risk assessment from scratch with clean table structure.
    5 rows × 2 cols, compact spacing. Labels: 楷体 bold. Values: 仿宋.
    """
    from docx.shared import Pt as Pt2

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt2(10.5)
    style.paragraph_format.space_before = Pt2(0)
    style.paragraph_format.space_after = Pt2(0)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("风险评估报备表")
    _set_font(title_run, FONT_TITLE, 22, bold=True)

    # Header lines
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("填报单位（盖章）：{{ reporting_unit }}         时间：{{ report_date }}")
    _set_font(h1_run, FONT_BODY, 10.5)

    h2 = doc.add_paragraph()
    h2_run = h2.add_run("联系人：{{ contact_person }}         联系电话：{{ contact_phone }}")
    _set_font(h2_run, FONT_BODY, 10.5)

    # Table — 5 rows × 2 cols
    table = _add_bordered_table(doc, 5, 2)

    # Row 0: 项目名称
    _add_cell_text(table.cell(0, 0), "项目名称", FONT_LABEL, 10.5, bold=True)
    _add_cell_text(table.cell(0, 1), "{{ project_name }}")
    _reduce_cell_spacing(table.cell(0, 1))

    # Row 1: 项目
    _add_cell_text(table.cell(1, 0), "项目", FONT_LABEL, 10.5, bold=True)
    _add_cell_text(table.cell(1, 1), "重大活动    {{ activity_type }}")
    _reduce_cell_spacing(table.cell(1, 1))

    # Row 2: 项目简要情况
    _add_cell_text(table.cell(2, 0), "项目简要情况", FONT_LABEL, 10.5, bold=True)
    detail = (
        "主办方：{{ sponsor }}\n"
        "承办方：{{ organizer }}\n"
        "活动参与方：{{ participants }}\n"
        "活动时间：{{ activity_start }} 至 {{ activity_end }}\n"
        "活动地点：{{ activity_location }}\n"
        "室内/户外：{{ is_indoor }}    场所类型：{{ location_type }}\n"
        "活动内容：{{ activity_content }}\n"
        "预计参与人数规模：{{ crowd_scale }}\n"
        "工作人员：{{ staff_count }}    安保人员：{{ security_count }}\n"
        "门票销售：{{ has_tickets }}\n"
        "是否有媒体：{{ has_media }}"
        "{% if has_media == '是' %}"
        "    采录方式：{{ media_channel }}    媒体名称：{{ media_name }}（{{ media_type }}）"
        "{% endif %}"
    )
    _replace_cell_text(table.cell(2, 1), detail)
    _reduce_cell_spacing(table.cell(2, 1))

    # Row 3: 主要风险因素
    _add_cell_text(table.cell(3, 0), "主要风险因素", FONT_LABEL, 10.5, bold=True)
    _replace_cell_text(table.cell(3, 1),
                       "{% for rf in risk_factors %}{{ loop.index }}. {{ rf }}\n{% endfor %}")
    _reduce_cell_spacing(table.cell(3, 1))

    # Row 4: 防范化解措施
    _add_cell_text(table.cell(4, 0), "防范化解措施", FONT_LABEL, 10.5, bold=True)
    _replace_cell_text(table.cell(4, 1),
                       "{% for mm in mitigation_measures %}{{ loop.index }}. {{ mm }}\n{% endfor %}")
    _reduce_cell_spacing(table.cell(4, 1))

    # Signature section
    for text in [
        "评估主体负责人签字：",
        "{{ assessor_signature }}",
        "安保负责人审核签字：",
        "{{ manager_signature }}",
    ]:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sr = sp.add_run(text)
        _set_font(sr, FONT_BODY, 10.5)

    output_path = TEMPLATES_ROOT / "risk_assessment" / "template.docx"
    doc.save(str(output_path))
    print(f"  ✅ risk_assessment → {output_path}")


# ---------------------------------------------------------------------------
# responsibility_letter — edit converted source (paragraph-based)
# ---------------------------------------------------------------------------
def build_responsibility_letter():
    """Edit converted 主办单位安全和消防责任确认书.docx. Replace with placeholders, set fonts."""
    doc = Document(str(SRC_DIR / "主办单位安全和消防责任确认书.docx"))

    # The source has 14 paragraphs: title, preamble, 8 declarations, 公章, 签字, 日期, 地点
    paras = doc.paragraphs

    # Title
    _replace_para_text(paras[0], "主办单位安全和消防责任确认书",
                       font=FONT_TITLE, size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Preamble — static text
    _replace_para_text(paras[1],
                       "依据国务院《大型群众性活动安全管理条例》、《中华人民共和国安全生产法》、"
                       "《中华人民共和国消防法》等法律法规，我单位对在负责范围内举办活动的安全和消防责任，"
                       "依法确认如下：",
                       font=FONT_BODY, size=10.5)

    # 8 declarations — hardcoded statutory text
    declaration_texts = [
        "一、已具体制订活动安全工作方案和安全责任制度，明确各项安全措施、落实安全工作人员岗位职责，"
        "事先开展活动安全宣传教育；",
        "二、保证临时搭建和使用的设施安全，沒有安全隐患；",
        "三、已按照负责许可的公安机关的要求，配备必要的安全检查设备，对参加活动的人员进行安全检查，"
        "对拒不接受安全检查的，有权拒绝进入；",
        "四、严格按照核准的活动场所容纳人员数量、划定的区域组织活动；",
        "五、已落实医疗救护、灭火、应急疏散等应急救援措施并组织演练；",
        "六、对妨碍活动安全的行为及时予以制止，发现违法犯罪行为及时向公安机关报告；",
        "七、已配备与活动安全工作需要相适应的专业保安人员或其他安全工作人员；",
        "八、已为活动的安全工作提供必要的保障。临时用电按规范采取安全措施。",
    ]

    for i, decl_text in enumerate(declaration_texts):
        para_idx = 2 + i  # paragraphs 2-9
        _replace_para_text(paras[para_idx], decl_text, font=FONT_BODY, size=10.5)

    # Signature block — right-aligned: 公章, 签字, signature image, 日期, 地点
    # 公章 (para 10)
    _replace_para_text(paras[10],
                       "活动主办单位（公章）：{{ sponsor_unit }}",
                       font=FONT_BODY, size=10.5, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # 签字 (para 11)
    _replace_para_text(paras[11],
                       "活动安全负责人（签字）：{{ security_leader_name }}",
                       font=FONT_BODY, size=10.5, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Signature image — below 签字 line, also right-aligned
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_p.add_run("{{ security_leader_signature }}")
    _set_font(sig_run, FONT_BODY, 10.5)

    # 日期 (para 12)
    _replace_para_text(paras[12],
                       "确认日期：{{ confirm_date }}",
                       font=FONT_BODY, size=10.5, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # 地点 (para 13)
    _replace_para_text(paras[13],
                       "确认地点：{{ confirm_location }}",
                       font=FONT_BODY, size=10.5, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    output_path = TEMPLATES_ROOT / "responsibility_letter" / "template.docx"
    doc.save(str(output_path))
    print(f"  ✅ responsibility_letter → {output_path}")


# ---------------------------------------------------------------------------
# filing_commitment — build from scratch with proper formatting
# ---------------------------------------------------------------------------
def build_filing_commitment():
    """Build filing commitment from scratch.
    Body paragraphs: first-line indent 2 chars, 仿宋_GB2312 12pt.
    Signature block: right-aligned.
    """
    from docx.shared import Pt as Pt2, Cm as Cm2

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt2(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt2(0)
    style.paragraph_format.space_after = Pt2(0)

    def add_body_para(text):
        """Add a body paragraph with first-line indent 2 chars."""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm2(0.85)
        r = p.add_run(text)
        _set_font(r, FONT_BODY, 12)
        return p

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("活动备案承诺书")
    _set_font(title_run, FONT_TITLE, 22, bold=True)

    # Addressee
    add_body_para("致天津市公安局：")

    # Activity info paragraph
    add_body_para(
        '我单位（{{ sponsor }}）定于{{ estimated_time }}，'
        '在{{ location }}举办"{{ project_name }}"活动。'
        '活动类型：{{ activity_type }}，预计参加人数：{{ crowd_scale }}，'
        '安保人员配置：{{ security_staff_count }}人。'
    )

    # Commitment header
    add_body_para("我单位郑重承诺：")

    # 4 commitment clauses
    for clause in [
        "一、所提交的活动方案、安保方案、风险评估报备表、安全消防责任确认书等备案材料真实、完整、有效；",
        "二、严格遵守《大型群众性活动安全管理条例》及相关法律法规，落实安全主体责任；",
        "三、按照批准的安保方案组织实施，确保活动安全有序进行；",
        "四、如活动时间、地点、规模等发生变更，及时向公安机关重新报备。",
    ]:
        add_body_para(clause)

    # Liability statement
    add_body_para("以上承诺如有虚假，我单位愿承担相应法律责任。")

    # Closing
    add_body_para("特此承诺")

    # One empty line before signature block
    sig_gap = doc.add_paragraph()
    sig_gap.paragraph_format.space_before = Pt2(0)
    sig_gap.paragraph_format.space_after = Pt2(0)

    # Right-aligned signature block
    for sf in [
        "承诺单位（盖章）：{{ sponsor }}",
        "安全负责人签字：{{ manager_signature }}",
        "日期：{{ filing_date }}",
    ]:
        sig_p = doc.add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_p.add_run(sf)
        _set_font(sig_run, FONT_BODY, 12)

    output_path = TEMPLATES_ROOT / "filing_commitment" / "template.docx"
    doc.save(str(output_path))
    print(f"  ✅ filing_commitment → {output_path}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _replace_para_text(para, new_text, font=FONT_BODY, size=10.5, bold=False, alignment=None):
    """Replace all text in a paragraph with new text and set font."""
    # Clear existing runs
    for r in para.runs:
        r.text = ''
    # If paragraph is empty, add a run
    if not para.runs:
        para.add_run('')
    para.runs[0].text = new_text
    _set_font(para.runs[0], font, size, bold)
    if alignment is not None:
        para.alignment = alignment
    # Clear other runs
    for r in para.runs[1:]:
        r.text = ''


def _replace_cell_text(cell, new_text, font=FONT_BODY, size=10.5, bold=False):
    """Replace all text in a table cell and set font."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run('')
    p.runs[0].text = new_text
    _set_font(p.runs[0], font, size, bold)
    # Clear other runs
    for r in p.runs[1:]:
        r.text = ''


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("Rebuilding all 5 template DOCX files...\n")
    build_activity_plan()
    build_security_plan()
    build_risk_assessment()
    build_responsibility_letter()
    build_filing_commitment()
    print("\nDone. All templates rebuilt.")


if __name__ == "__main__":
    main()
