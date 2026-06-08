from __future__ import annotations

import asyncio
import hashlib
import logging
import subprocess
import tempfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

from docxtpl import DocxTemplate
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.activity import Activity, ActivityPlan, SecurityPlan
from app.models.material import KeyMaterial, SecurityPlanMaterial
from app.models.template import FilledDocument
from app.models.document import Document
from app.services import minio_client
from app.templates import (
    FORM_MODELS, SCHEMAS, TEMPLATES_ROOT, TEMPLATE_DISPLAY_NAMES, TEMPLATE_ENTITY_MAP,
)
from app.templates.security_plan.schema import CONDITIONAL_FIELDS

logger = logging.getLogger("camis.template")

MINIO_BUCKET = "camis2026"


class TemplateService:
    def __init__(self, db: AsyncSession):
        self.db = db

    # ------------------------------------------------------------------
    # material lifecycle
    # ------------------------------------------------------------------

    async def get_or_create_material(self, activity_id: UUID, material_type: str) -> KeyMaterial:
        """Get existing KeyMaterial by (activity_id, material_type) or create one."""
        result = await self.db.execute(
            select(KeyMaterial).where(
                KeyMaterial.activity_id == activity_id,
                KeyMaterial.material_type == material_type,
            )
        )
        entity = result.scalar_one_or_none()
        if entity:
            return entity
        display_name = TEMPLATE_DISPLAY_NAMES.get(material_type, material_type)
        entity = KeyMaterial(name=display_name, activity_id=activity_id, material_type=material_type)
        self.db.add(entity)
        await self.db.flush()
        sp = await self.db.execute(
            select(SecurityPlan).where(SecurityPlan.activity_id == activity_id)
        )
        sp = sp.scalar_one_or_none()
        if sp:
            self.db.add(SecurityPlanMaterial(security_plan_id=sp.id, material_id=entity.id))
            await self.db.flush()
        await self.db.commit()
        return entity

    # ------------------------------------------------------------------
    # schema + draft
    # ------------------------------------------------------------------

    async def get_schema(
        self, template_type: str, activity_id: UUID, material_type: str | None = None,
    ) -> dict:
        schema = SCHEMAS[template_type].copy()
        schema["template_type"] = template_type

        entity = await self._get_entity(template_type, activity_id, material_type)
        draft_data = getattr(entity, "draft_data", None) if entity else None
        current_fd_id = getattr(entity, "current_filled_document_id", None) if entity else None

        schema["has_draft"] = draft_data is not None
        schema["draft_data"] = draft_data

        # latest version snapshot (pre-fill when no draft)
        current_version = None
        snapshot_data = None
        if current_fd_id:
            fd = await self.db.get(FilledDocument, current_fd_id)
            if fd:
                current_version = fd.version_number
                snapshot_data = fd.data_snapshot
        schema["current_version"] = current_version
        schema["snapshot_data"] = snapshot_data

        # security plan: expose risk_level for conditional field filtering
        if template_type == "security_plan":
            risk_level = getattr(entity, "risk_level", None) if entity else None
            schema["risk_level"] = risk_level

        # autofill: provide Activity + ActivityPlan snapshot + SecurityPlan snapshot + defaults
        ACTIVITY_FIELD_MAP = {
            "project_name": "name",
            "sponsor": "sponsor",
            "location_type": "location",
            "activity_type": "type",
            "location": "location",
            "estimated_time": "estimated_time",
        }
        PLAN_FIELD_MAP = {
            "activity_content": "activity_content",
            "activity_start": "start_time",
            "activity_end": "end_time",
            "staff_count": "staff_count",
            "crowd_scale": "opening_crowd",
        }
        SECURITY_PLAN_FIELD_MAP = {
            "security_staff_count": "security_staff_count",
            "security_count": "security_staff_count",
        }
        # fixed defaults (sensitive values from .env via settings)
        from app.config import settings
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        DEFAULT_FIELD_MAP: dict[str, str] = {
            "reporting_unit": settings.default_reporting_unit or "",
            "report_date": today,
            "sponsor_unit": settings.default_reporting_unit or "",
            "confirm_date": today,
            "confirm_location": settings.default_confirm_location or "",
        }
        all_field_names = {f["name"] for f in schema.get("fields", [])}
        mappable = all_field_names & (set(ACTIVITY_FIELD_MAP) | set(PLAN_FIELD_MAP) | set(SECURITY_PLAN_FIELD_MAP) | set(DEFAULT_FIELD_MAP))
        if mappable:
            autofill_data: dict[str, object] = {}
            activity = await self.db.get(Activity, activity_id)
            plan_snapshot: dict[str, object] = {}
            plan_result = await self.db.execute(
                select(ActivityPlan).where(ActivityPlan.activity_id == activity_id)
            )
            plan = plan_result.scalar_one_or_none()
            if plan:
                if plan.current_filled_document_id:
                    plan_fd = await self.db.get(FilledDocument, plan.current_filled_document_id)
                    if plan_fd and plan_fd.data_snapshot:
                        plan_snapshot = dict(plan_fd.data_snapshot)
                if plan.draft_data:
                    plan_snapshot.update(plan.draft_data)
            sec_snapshot: dict[str, object] = {}
            if mappable & set(SECURITY_PLAN_FIELD_MAP):
                sec_result = await self.db.execute(
                    select(SecurityPlan).where(SecurityPlan.activity_id == activity_id)
                )
                sec = sec_result.scalar_one_or_none()
                if sec:
                    if sec.current_filled_document_id:
                        sec_fd = await self.db.get(FilledDocument, sec.current_filled_document_id)
                        if sec_fd and sec_fd.data_snapshot:
                            sec_snapshot = dict(sec_fd.data_snapshot)
                    if sec.draft_data:
                        sec_snapshot.update(sec.draft_data)
            for name in mappable:
                if name in ACTIVITY_FIELD_MAP and activity:
                    attr = ACTIVITY_FIELD_MAP[name]
                    val = getattr(activity, attr, None)
                    if isinstance(val, datetime):
                        val = val.strftime("%Y年%m月%d日")
                    autofill_data[name] = val or ""
                elif name in PLAN_FIELD_MAP:
                    plan_key = PLAN_FIELD_MAP[name]
                    val = plan_snapshot.get(plan_key, "") or ""
                    if plan_key == "opening_crowd" and not val:
                        val = plan_snapshot.get("regular_crowd", "") or ""
                    autofill_data[name] = val
                elif name in SECURITY_PLAN_FIELD_MAP:
                    sec_key = SECURITY_PLAN_FIELD_MAP[name]
                    autofill_data[name] = sec_snapshot.get(sec_key, "") or ""
                elif name in DEFAULT_FIELD_MAP:
                    autofill_data[name] = DEFAULT_FIELD_MAP[name]
            schema["autofill_data"] = autofill_data

        return schema

    async def save_draft(
        self, template_type: str, activity_id: UUID, data: dict, user_id: UUID,
        material_type: str | None = None,
    ) -> None:
        entity = await self._get_or_create_entity(template_type, activity_id, user_id, material_type)
        entity.draft_data = data
        await self.db.flush()
        await self.db.commit()
        logger.info("draft saved type=%s activity=%s", template_type, activity_id)

    # ------------------------------------------------------------------
    # generate
    # ------------------------------------------------------------------

    async def generate(
        self, template_type: str, activity_id: UUID, data: dict, user_id: UUID,
        material_type: str | None = None,
    ) -> dict:
        entity = await self._get_or_create_entity(template_type, activity_id, user_id, material_type)

        # resolve version
        version_number = await self._next_version(activity_id, template_type)
        risk_level = getattr(entity, "risk_level", None)

        # DOCX deferred for types that require Manager signing
        DEFERRED_TYPES = {"security_plan", "risk_assessment", "responsibility_letter", "filing_commitment"}
        is_deferred = template_type in DEFERRED_TYPES

        docx_bytes = None
        minio_path = None
        template_hash = ""

        if not is_deferred:
            docx_bytes = await self._render_docx(template_type, data, risk_level)
            template_hash = hashlib.sha256(
                (TEMPLATES_ROOT / template_type / "template.docx").read_bytes()
            ).hexdigest()
            minio_path = f"filled_documents/{activity_id}/{template_type}/v{version_number}.docx"
            await minio_client.upload_file(minio_path, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # create FilledDocument
        fd = FilledDocument(
            activity_id=activity_id,
            template_type=template_type,
            version_number=version_number,
            data_snapshot=data,
            minio_path=minio_path,
            pdf_path=None,
            template_hash=template_hash,
            generated_by=user_id,
        )
        self.db.add(fd)
        await self.db.flush()

        if not is_deferred and docx_bytes:
            asyncio.create_task(render_pdf_background(fd.id, docx_bytes, activity_id, template_type, version_number))

        # link to entity
        entity.current_filled_document_id = fd.id
        entity.draft_data = None
        entity.submit_time = datetime.now(timezone.utc)
        await self.db.flush()

        # also link KeyMaterial so list_materials can find the minio_path for preview
        km_result = await self.db.execute(
            select(KeyMaterial).where(
                KeyMaterial.activity_id == activity_id,
                KeyMaterial.material_type == template_type,
            )
        )
        km = km_result.scalar_one_or_none()
        if km:
            km.current_filled_document_id = fd.id
            await self.db.flush()

        # cross-template sync: security_staff_count → risk_assessment.security_count
        if template_type == "security_plan" and "security_staff_count" in data:
            for mt in ["risk_assessment", "filing_commitment"]:
                mt_km_result = await self.db.execute(
                    select(KeyMaterial).where(
                        KeyMaterial.activity_id == activity_id,
                        KeyMaterial.material_type == mt,
                    )
                )
                mt_km = mt_km_result.scalar_one_or_none()
                if mt_km and mt_km.current_filled_document_id:
                    mt_fd = await self.db.get(FilledDocument, mt_km.current_filled_document_id)
                    if mt_fd and mt_fd.data_snapshot:
                        mt_data = dict(mt_fd.data_snapshot)
                        field_name = "security_count" if mt == "risk_assessment" else "security_staff_count"
                        mt_data[field_name] = data["security_staff_count"]
                        mt_vn = await self._next_version(activity_id, mt)
                        mt_new_fd = FilledDocument(
                            activity_id=activity_id, template_type=mt,
                            version_number=mt_vn, data_snapshot=mt_data,
                            minio_path=None,  # deferred
                            generated_by=user_id,
                        )
                        self.db.add(mt_new_fd)
                        await self.db.flush()
                        mt_km.current_filled_document_id = mt_new_fd.id
                        mt_km.draft_data = None
                        logger.info("cross-sync: %s v%d updated from security_plan", mt, mt_vn)

        await self.db.commit()

        logger.info(
            "generated type=%s activity=%s v%d deferred=%s",
            template_type, activity_id, version_number, is_deferred,
        )
        return {
            "id": fd.id,
            "template_type": fd.template_type,
            "version_number": fd.version_number,
            "minio_path": fd.minio_path,
            "pdf_ready": False,
            "pdf_preview_url": None,
            "docx_bytes": docx_bytes,
            "created_at": fd.created_at.isoformat() if fd.created_at else None,
        }

    # ------------------------------------------------------------------
    # finalize
    # ------------------------------------------------------------------

    async def finalize_plan(self, activity_id: UUID, user_id: UUID) -> None:
        """Finalize the activity plan: validate content, trigger workflow transition."""
        entity = await self._get_entity("activity_plan", activity_id)
        if not entity or not entity.current_filled_document_id:
            raise ValueError("活动方案尚未生成，无法最终确定")

        fd = await self.db.get(FilledDocument, entity.current_filled_document_id)
        if not fd or not fd.data_snapshot:
            raise ValueError("未找到当前版本数据")

        from app.templates.activity_plan.schema import ActivityPlanForm, SCHEMA as PLAN_SCHEMA

        # clean snapshot: strip hidden conditional fields, fix int→str for select fields
        cleaned = dict(fd.data_snapshot)
        for f in PLAN_SCHEMA["fields"]:
            cond = f.get("condition")
            if cond:
                parts = cond.replace("==", " ").replace("!=", " ").split()
                if len(parts) >= 3:
                    cur = cleaned.get(parts[0], "")
                    target = parts[2].strip("'\"")
                    match = cur == target if "==" in cond else cur != target
                    if not match:
                        cleaned.pop(f["name"], None)
        select_fields = {f["name"] for f in PLAN_SCHEMA["fields"] if f.get("ui_type") == "select"}
        for name in select_fields:
            if isinstance(cleaned.get(name), int):
                cleaned[name] = str(cleaned[name]) if cleaned[name] != 0 else ""

        try:
            ActivityPlanForm(**cleaned)
        except Exception as e:
            raise ValueError(f"活动方案内容不完整: {e}")

        data = fd.data_snapshot
        import re

        errors: list[str] = []
        if not data.get("activity_content"):
            errors.append("活动主要内容不能为空")
        if not data.get("start_time"):
            errors.append("开始时间未填写")
        if not data.get("end_time"):
            errors.append("结束时间未填写")
        if data.get("start_time") and data.get("end_time") and data["start_time"] >= data["end_time"]:
            errors.append("结束时间必须晚于开始时间")
        if not data.get("staff_count") or data["staff_count"] <= 0:
            errors.append("工作人员数量必须大于0")
        if not data.get("construction_plan"):
            errors.append("搭建方案不能为空")
        if not data.get("regular_crowd"):
            errors.append("平日人数范围未选择")
        phone = data.get("contact_phone", "")
        if not re.match(r"^1[3-9]\d{9}$", str(phone)):
            errors.append("负责人联系方式须为11位手机号码")
        if data.get("has_opening") == "是":
            if not data.get("opening_start"):
                errors.append("开幕式开始时间未填写")
            if not data.get("opening_end"):
                errors.append("开幕式结束时间未填写")
            if not data.get("opening_crowd"):
                errors.append("主要活动日人数范围未选择")
        if data.get("has_performers") == "是":
            if not data.get("performer_count") or data["performer_count"] <= 0:
                errors.append("演员数量必须大于0")
            if not data.get("guest_count") or data["guest_count"] <= 0:
                errors.append("嘉宾数量必须大于0")

        if errors:
            raise ValueError("; ".join(errors))

        activity = await self.db.get(Activity, activity_id)
        if not activity or activity.status != "待设计方案":
            raise ValueError("当前状态不允许最终确定")

        from app.services.workflow_service import WorkflowService
        from app.services.notification_service import NotificationService
        ws = WorkflowService(self.db, NotificationService(self.db))
        User = __import__("app.models.user", fromlist=["User"]).User
        await ws.transition(activity_id, "待安保方案设计", await self.db.get(User, user_id))

    async def submit_security_plan_for_review(self, activity_id: UUID, user_id: UUID) -> None:
        """Submit security plan for SecurityManager review. Validate content, set audit_status=待签署."""
        entity = await self._get_entity("security_plan", activity_id)
        if not entity or not entity.current_filled_document_id:
            raise ValueError("安保方案尚未生成，无法提交审核")

        if not getattr(entity, "risk_level", None):
            raise ValueError("请先选择风险等级")

        fd = await self.db.get(FilledDocument, entity.current_filled_document_id)
        if not fd or not fd.data_snapshot:
            raise ValueError("未找到当前版本数据")

        from app.templates.security_plan.schema import SecurityPlanForm, CONDITIONAL_FIELDS as SP_CONDITIONAL
        import re

        # strip conditional fields not applicable to current risk_level
        cleaned = dict(fd.data_snapshot)
        risk_level = getattr(entity, "risk_level", "") or ""
        allowed = set(SP_CONDITIONAL.get(risk_level, []))
        for cond_field in {"medical_plan", "fire_plan", "crowd_control"}:
            if cond_field not in allowed:
                cleaned.pop(cond_field, None)

        try:
            SecurityPlanForm(**cleaned)
        except Exception as e:
            raise ValueError(f"安保方案内容不完整: {e}")

        data = fd.data_snapshot
        errors: list[str] = []
        if not data.get("security_staff_config"):
            errors.append("安保人员配置不能为空")
        if not data.get("movement_plan"):
            errors.append("动线设计不能为空")
        if not data.get("equipment_list"):
            errors.append("安保设备清单不能为空")
        if not data.get("emergency_plan"):
            errors.append("应急预案不能为空")
        if not data.get("security_staff_count") or data["security_staff_count"] <= 0:
            errors.append("安保人员数量必须大于0")

        risk_level = getattr(entity, "risk_level", "") or ""
        if risk_level == "高风险" and not data.get("medical_plan"):
            errors.append("医疗救护措施不能为空（风险等级：高风险）")
        if risk_level in ("高风险", "中低风险") and not data.get("fire_plan"):
            errors.append("消防措施不能为空")
        if risk_level == "高风险" and not data.get("crowd_control"):
            errors.append("人流管控方案不能为空")

        if errors:
            raise ValueError("; ".join(errors))

        if entity.rejected_at and fd.created_at and fd.created_at <= entity.rejected_at:
            raise ValueError("驳回后请先生成新版本再提交审核")

        activity = await self.db.get(Activity, activity_id)
        if not activity or activity.status != "待安保方案设计":
            raise ValueError("当前状态不允许提交审核")

        entity.audit_status = "待签署"
        entity.last_reject_reason = None
        entity.rejected_at = None
        await self.db.commit()

        from app.services.notification_service import NotificationService
        ns = NotificationService(self.db)
        await ns.notify_role("SecurityManager", "安保方案已提交，请审核签署",
            reference_id=activity_id, reference_type="activity")

    async def sign_and_finalize(self, activity_id: UUID, user_id: UUID,
                                manager_signature: str) -> None:
        """SecurityManager step 1: sign 3 files (security_plan + 双表), generate DOCX.

        Does NOT transition — waits for sign_manager_commitment (step 2)."""
        from app.models.rbac import Role, UserRole

        role_rows = await self.db.execute(
            select(Role.name).join(UserRole, UserRole.role_id == Role.id)
            .where(UserRole.user_id == user_id)
        )
        roles = {row[0] for row in role_rows.all()}
        if "SecurityManager" not in roles:
            raise ValueError("仅安保负责人可签署确认")

        entity = await self._get_entity("security_plan", activity_id)
        if not entity or entity.audit_status != "待签署":
            raise ValueError("当前状态不允许签署")

        activity = await self.db.get(Activity, activity_id)
        if not activity or activity.status != "待安保方案设计":
            raise ValueError("当前活动状态不允许签署")

        # Generate DOCX for the 3 deferred types (NOT filing_commitment)
        SIGN_TYPES = ["security_plan", "risk_assessment", "responsibility_letter"]
        for ttype in SIGN_TYPES:
            fds = await self.db.execute(
                select(FilledDocument).where(
                    FilledDocument.activity_id == activity_id,
                    FilledDocument.template_type == ttype,
                    FilledDocument.minio_path.is_(None),
                )
            )
            for fd in fds.scalars().all():
                data = dict(fd.data_snapshot or {})
                data["manager_signature"] = manager_signature
                fd.data_snapshot = data

                risk_level = getattr(entity, "risk_level", None)
                docx_bytes = await self._render_docx(ttype, data, risk_level)
                fd.template_hash = hashlib.sha256(
                    (TEMPLATES_ROOT / ttype / "template.docx").read_bytes()
                ).hexdigest()
                minio_path = f"filled_documents/{activity_id}/{ttype}/v{fd.version_number}.docx"
                await minio_client.upload_file(minio_path, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                fd.minio_path = minio_path
                asyncio.create_task(render_pdf_background(fd.id, docx_bytes, activity_id, ttype, fd.version_number))

            # Ensure KeyMaterial exists, link FilledDocument, mark as signed
            km = await self.get_or_create_material(activity_id, ttype)
            t_entity = entity if ttype == "security_plan" else await self._get_entity(ttype, activity_id, ttype)
            if t_entity and t_entity.current_filled_document_id:
                km.current_filled_document_id = t_entity.current_filled_document_id
            km.sign_status = "signed"

        # Ensure activity_plan KeyMaterial exists, link its FilledDocument, mark as signed
        ap_km = await self.get_or_create_material(activity_id, "activity_plan")
        ap_result = await self.db.execute(
            select(ActivityPlan).where(ActivityPlan.activity_id == activity_id)
        )
        ap = ap_result.scalar_one_or_none()
        if ap and ap.current_filled_document_id:
            ap_km.current_filled_document_id = ap.current_filled_document_id
        ap_km.sign_status = "signed"

        # Update SecurityPlan
        entity.audit_status = "已签署"
        entity.manager_id = user_id
        entity.sign_time = datetime.now(timezone.utc)
        await self.db.commit()

    async def sign_manager_commitment(self, activity_id: UUID, user_id: UUID) -> None:
        """SecurityManager step 2: sign filing commitment, generate DOCX, transition to 待备案申请."""
        from app.models.rbac import Role, UserRole
        from app.services.workflow_service import WorkflowService
        from app.services.notification_service import NotificationService

        role_rows = await self.db.execute(
            select(Role.name).join(UserRole, UserRole.role_id == Role.id)
            .where(UserRole.user_id == user_id)
        )
        roles = {row[0] for row in role_rows.all()}
        if "SecurityManager" not in roles:
            raise ValueError("仅安保负责人可签署确认")

        entity = await self._get_entity("security_plan", activity_id)
        if not entity or entity.audit_status != "已签署":
            raise ValueError("请先签署安保方案、风险评估表和责任确认书")

        activity = await self.db.get(Activity, activity_id)
        if not activity or activity.status != "待安保方案设计":
            raise ValueError("当前活动状态不允许签署")

        # Get manager signature from any signed FilledDocument
        mgr_sig = ""
        for ttype in ["security_plan", "risk_assessment", "responsibility_letter"]:
            fd_result = await self.db.execute(
                select(FilledDocument).where(
                    FilledDocument.activity_id == activity_id,
                    FilledDocument.template_type == ttype,
                ).order_by(FilledDocument.version_number.desc()).limit(1)
            )
            fd = fd_result.scalar_one_or_none()
            if fd and fd.data_snapshot:
                mgr_sig = fd.data_snapshot.get("manager_signature", "")
                if mgr_sig:
                    break

        # Build autofill data for filing_commitment
        plan_result = await self.db.execute(
            select(ActivityPlan).where(ActivityPlan.activity_id == activity_id)
        )
        plan = plan_result.scalar_one_or_none()
        plan_snapshot: dict = {}
        if plan and plan.current_filled_document_id:
            plan_fd = await self.db.get(FilledDocument, plan.current_filled_document_id)
            if plan_fd and plan_fd.data_snapshot:
                plan_snapshot = plan_fd.data_snapshot

        sec_snapshot: dict = {}
        if entity.current_filled_document_id:
            sec_fd = await self.db.get(FilledDocument, entity.current_filled_document_id)
            if sec_fd and sec_fd.data_snapshot:
                sec_snapshot = sec_fd.data_snapshot

        commit_data = {
            "project_name": activity.name,
            "sponsor": activity.sponsor,
            "estimated_time": activity.estimated_time.strftime("%Y年%m月%d日") if activity.estimated_time else "",
            "location": activity.location,
            "activity_type": activity.type,
            "crowd_scale": plan_snapshot.get("opening_crowd") or plan_snapshot.get("regular_crowd", ""),
            "security_staff_count": str(sec_snapshot.get("security_staff_count", "")),
            "filing_date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "manager_signature": mgr_sig,
        }

        # Generate filing_commitment DOCX (create FilledDocument if not deferred)
        fds = await self.db.execute(
            select(FilledDocument).where(
                FilledDocument.activity_id == activity_id,
                FilledDocument.template_type == "filing_commitment",
                FilledDocument.minio_path.is_(None),
            )
        )
        fd_list = fds.scalars().all()
        commit_fd = None
        if fd_list:
            for fd in fd_list:
                fd.data_snapshot = {**fd.data_snapshot, "manager_signature": mgr_sig}
                docx_bytes = await self._render_docx("filing_commitment", dict(fd.data_snapshot), None)
                fd.template_hash = hashlib.sha256(
                    (TEMPLATES_ROOT / "filing_commitment" / "template.docx").read_bytes()
                ).hexdigest()
                minio_path = f"filled_documents/{activity_id}/filing_commitment/v{fd.version_number}.docx"
                await minio_client.upload_file(minio_path, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                fd.minio_path = minio_path
                asyncio.create_task(render_pdf_background(fd.id, docx_bytes, activity_id, "filing_commitment", fd.version_number))
                commit_fd = fd
        else:
            # No deferred FilledDocument — create one on the fly
            vn = await self._next_version(activity_id, "filing_commitment")
            docx_bytes = await self._render_docx("filing_commitment", commit_data, None)
            template_hash = hashlib.sha256(
                (TEMPLATES_ROOT / "filing_commitment" / "template.docx").read_bytes()
            ).hexdigest()
            minio_path = f"filled_documents/{activity_id}/filing_commitment/v{vn}.docx"
            await minio_client.upload_file(minio_path, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            commit_fd = FilledDocument(
                activity_id=activity_id, template_type="filing_commitment",
                version_number=vn, data_snapshot=commit_data,
                minio_path=minio_path, template_hash=template_hash,
                generated_by=user_id,
            )
            self.db.add(commit_fd)
            await self.db.flush()
            asyncio.create_task(render_pdf_background(commit_fd.id, docx_bytes, activity_id, "filing_commitment", vn))

        # Ensure filing_commitment KeyMaterial exists, linked, and signed
        km = await self.get_or_create_material(activity_id, "filing_commitment")
        km.current_filled_document_id = commit_fd.id
        km.sign_status = "signed"

        await self.db.commit()

        # Transition workflow
        ws = WorkflowService(self.db, NotificationService(self.db))
        User = __import__("app.models.user", fromlist=["User"]).User
        await ws.transition(activity_id, "待备案申请", await self.db.get(User, user_id))

    async def reject_security_plan(self, activity_id: UUID, user_id: UUID,
                                   reasons: list[str], comment: str | None = None) -> None:
        """SecurityManager rejects the security plan back to SecurityOfficer."""
        from app.models.rbac import Role, UserRole
        from app.services.notification_service import NotificationService

        role_rows = await self.db.execute(
            select(Role.name).join(UserRole, UserRole.role_id == Role.id)
            .where(UserRole.user_id == user_id)
        )
        roles = {row[0] for row in role_rows.all()}
        if "SecurityManager" not in roles:
            raise ValueError("仅安保负责人可驳回")

        entity = await self._get_entity("security_plan", activity_id)
        if not entity or entity.audit_status != "待签署":
            raise ValueError("当前状态不允许驳回")

        full_reason = "；".join(reasons)
        if comment:
            full_reason += f"（补充：{comment}）"

        entity.audit_status = "待编制"
        entity.last_reject_reason = full_reason
        entity.rejected_at = datetime.now(timezone.utc)
        entity.reject_count = (entity.reject_count or 0) + 1
        await self.db.commit()

        ns = NotificationService(self.db)
        await ns.notify_role("SecurityOfficer", f"安保方案被驳回需修改：{full_reason}",
            reference_id=activity_id, reference_type="activity")

    # ------------------------------------------------------------------
    # versions
    # ------------------------------------------------------------------

    async def get_versions(
        self, template_type: str, activity_id: UUID, material_type: str | None = None,
    ) -> list[dict]:
        result = await self.db.execute(
            select(FilledDocument)
            .where(
                FilledDocument.activity_id == activity_id,
                FilledDocument.template_type == template_type,
            )
            .order_by(FilledDocument.version_number.desc())
        )
        rows = result.scalars().all()

        entity = await self._get_entity(template_type, activity_id, material_type)
        current_id = getattr(entity, "current_filled_document_id", None) if entity else None

        return [
            {
                "id": str(r.id),
                "version_number": r.version_number,
                "generated_by": str(r.generated_by),
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "is_current": r.id == current_id,
                "pdf_ready": r.pdf_path is not None,
            }
            for r in rows
        ]

    async def get_version_detail(
        self, template_type: str, activity_id: UUID, version_number: int,
        material_type: str | None = None,
    ) -> dict | None:
        result = await self.db.execute(
            select(FilledDocument).where(
                FilledDocument.activity_id == activity_id,
                FilledDocument.template_type == template_type,
                FilledDocument.version_number == version_number,
            )
        )
        fd = result.scalar_one_or_none()
        if not fd:
            return None

        entity = await self._get_entity(template_type, activity_id, material_type)
        current_id = getattr(entity, "current_filled_document_id", None) if entity else None

        return {
            "id": str(fd.id),
            "version_number": fd.version_number,
            "data_snapshot": fd.data_snapshot,
            "template_hash": fd.template_hash,
            "generated_by": str(fd.generated_by),
            "created_at": fd.created_at.isoformat() if fd.created_at else None,
            "is_current": fd.id == current_id,
            "pdf_ready": fd.pdf_path is not None,
        }

    async def get_version_preview_url(
        self, template_type: str, activity_id: UUID, version_number: int,
    ) -> str | None:
        fd = await self._get_version_row(activity_id, template_type, version_number)
        if not fd or not fd.pdf_path:
            return None
        return await minio_client.get_presigned_url(fd.pdf_path, inline=True)

    async def get_version_diff(
        self, template_type: str, activity_id: UUID, v1: int, v2: int,
        material_type: str | None = None,
    ) -> list[dict]:
        fd1 = await self._get_version_row(activity_id, template_type, v1)
        fd2 = await self._get_version_row(activity_id, template_type, v2)
        if not fd1 or not fd2:
            return []

        s1 = fd1.data_snapshot or {}
        s2 = fd2.data_snapshot or {}
        all_keys = set(s1) | set(s2)
        diffs = []
        for key in sorted(all_keys):
            old_val = s1.get(key)
            new_val = s2.get(key)
            if old_val != new_val:
                diffs.append({"field": key, "old": old_val, "new": new_val})
        return diffs

    # ------------------------------------------------------------------
    # helpers
    # ------------------------------------------------------------------

    async def _render_docx(
        self, template_type: str, data: dict, risk_level: str | None,
    ) -> bytes:
        from docxtpl import InlineImage
        from docx.shared import Mm

        template_path = TEMPLATES_ROOT / template_type / "template.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"template not found: {template_path}")

        doc = DocxTemplate(template_path)
        context = dict(data)

        # detect signature fields and embed images if value is a minio_path
        for field_name, value in list(context.items()):
            if "signature" in field_name and isinstance(value, str) and value:
                try:
                    img_bytes = minio_client.minio_client.get_object(
                        minio_client._bucket, value,
                    ).read()
                    context[field_name] = InlineImage(doc, BytesIO(img_bytes), width=Mm(30))
                except Exception:
                    pass  # fall back to text rendering

        doc.render(context)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def _next_version(self, activity_id: UUID, template_type: str) -> int:
        result = await self.db.execute(
            select(func.max(FilledDocument.version_number)).where(
                FilledDocument.activity_id == activity_id,
                FilledDocument.template_type == template_type,
            )
        )
        current_max = result.scalar() or 0
        return current_max + 1

    async def _get_version_row(
        self, activity_id: UUID, template_type: str, version_number: int,
    ) -> FilledDocument | None:
        result = await self.db.execute(
            select(FilledDocument).where(
                FilledDocument.activity_id == activity_id,
                FilledDocument.template_type == template_type,
                FilledDocument.version_number == version_number,
            )
        )
        return result.scalar_one_or_none()

    async def _get_entity(
        self, template_type: str, activity_id: UUID, material_type: str | None = None,
    ):
        entity_type = TEMPLATE_ENTITY_MAP[template_type]
        if entity_type == "activity_plan":
            result = await self.db.execute(
                select(ActivityPlan).where(ActivityPlan.activity_id == activity_id)
            )
            return result.scalar_one_or_none()
        elif entity_type == "security_plan":
            result = await self.db.execute(
                select(SecurityPlan).where(SecurityPlan.activity_id == activity_id)
            )
            return result.scalar_one_or_none()
        elif entity_type == "key_material":
            if not material_type:
                raise ValueError("material_type required for key_material templates")
            result = await self.db.execute(
                select(KeyMaterial).where(
                    KeyMaterial.activity_id == activity_id,
                    KeyMaterial.material_type == material_type,
                )
            )
            return result.scalar_one_or_none()
        return None

    async def _get_or_create_entity(
        self, template_type: str, activity_id: UUID, user_id: UUID,
        material_type: str | None = None,
    ):
        entity_type = TEMPLATE_ENTITY_MAP[template_type]
        if entity_type == "activity_plan":
            result = await self.db.execute(
                select(ActivityPlan).where(ActivityPlan.activity_id == activity_id)
            )
            entity = result.scalar_one_or_none()
            if not entity:
                entity = ActivityPlan(activity_id=activity_id, designer_id=user_id)
                self.db.add(entity)
                await self.db.flush()
            # Ensure KeyMaterial exists and material_id is backfilled
            if not entity.material_id:
                km = await self.get_or_create_material(activity_id, "activity_plan")
                entity.material_id = km.id
                await self.db.flush()
            return entity
        elif entity_type == "security_plan":
            result = await self.db.execute(
                select(SecurityPlan).where(SecurityPlan.activity_id == activity_id)
            )
            entity = result.scalar_one_or_none()
            if not entity:
                entity = SecurityPlan(activity_id=activity_id)
                self.db.add(entity)
                await self.db.flush()
            if not entity.material_id:
                km = await self.get_or_create_material(activity_id, "security_plan")
                entity.material_id = km.id
                await self.db.flush()
            return entity
        elif entity_type == "key_material":
            if not material_type:
                raise ValueError("material_type required for key_material templates")
            result = await self.db.execute(
                select(KeyMaterial).where(
                    KeyMaterial.activity_id == activity_id,
                    KeyMaterial.material_type == material_type,
                )
            )
            entity = result.scalar_one_or_none()
            if not entity:
                display_name = TEMPLATE_DISPLAY_NAMES.get(template_type, template_type)
                entity = KeyMaterial(
                    name=display_name,
                    activity_id=activity_id,
                    material_type=material_type,
                )
                self.db.add(entity)
                await self.db.flush()
                sp = await self.db.execute(
                    select(SecurityPlan).where(SecurityPlan.activity_id == activity_id)
                )
                sp = sp.scalar_one_or_none()
                if sp:
                    self.db.add(SecurityPlanMaterial(security_plan_id=sp.id, material_id=entity.id))
                    await self.db.flush()
            return entity
        raise ValueError(f"unknown entity type: {entity_type}")


# ── background PDF renderer (standalone, own DB session) ──

async def render_pdf_background(
    fd_id: UUID, docx_bytes: bytes, activity_id: UUID,
    template_type: str, version_number: int,
) -> None:
    """Generate PDF in background with its own DB session."""
    from app.database import async_session

    async with async_session() as db:
        pdf_path = f"filled_documents/{activity_id}/{template_type}/v{version_number}.pdf"
        try:
            pdf_bytes = await _docx_to_pdf_sync(docx_bytes)
            await minio_client.upload_file(pdf_path, pdf_bytes, "application/pdf")
            fd = await db.get(FilledDocument, fd_id)
            if fd:
                fd.pdf_path = pdf_path
                await db.commit()
            logger.info("pdf render ok fd=%s path=%s", fd_id, pdf_path)
        except Exception as e:
            logger.warning("pdf render failed fd=%s type=%s: %s", fd_id, template_type, e)
            import traceback
            traceback.print_exc()


async def _docx_to_pdf_sync(docx_bytes: bytes) -> bytes:
    """Convert DOCX to PDF via LibreOffice in a thread — must not block event loop."""
    import asyncio

    def _convert() -> bytes:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf_in:
            tf_in.write(docx_bytes)
            in_path = tf_in.name

        out_dir = tempfile.mkdtemp()
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path],
                check=True, timeout=60, capture_output=True,
            )
            pdf_name = Path(in_path).stem + ".pdf"
            pdf_path = Path(out_dir) / pdf_name
            if not pdf_path.exists():
                raise RuntimeError("pdf not produced by soffice")
            return pdf_path.read_bytes()
        finally:
            Path(in_path).unlink(missing_ok=True)
            import shutil
            shutil.rmtree(out_dir, ignore_errors=True)

    return await asyncio.to_thread(_convert)
